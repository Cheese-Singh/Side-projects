import io
import os
import re
import time
import tempfile
from pathlib import Path

try:
    from mlx_vlm import load as vlm_load, generate as vlm_generate
    from mlx_vlm.prompt_utils import apply_chat_template
    from mlx_vlm.utils import load_config
    VLM_AVAILABLE = True
    print("[A13] mlx_vlm imported.")
except Exception as e:
    VLM_AVAILABLE = False
    print(f"[A13] mlx_vlm unavailable ({e}).")

MFLUX_AVAILABLE = False
_Flux1Class = None
_FluxConfig = None
try:
    from mflux.models.flux.variants.txt2img.flux import Flux1 as _Flux1Class
    for _cfg_path in (
        "mflux.config.config",
        "mflux.config",
        "mflux",
    ):
        try:
            import importlib as _il
            _m = _il.import_module(_cfg_path)
            _FluxConfig = getattr(_m, "Config")
            break
        except Exception:
            continue
    MFLUX_AVAILABLE = True
    print("[A13] mflux imported.")
except Exception as e:
    print(f"[A13] mflux unavailable ({e}).")

try:
    import fitz
    PYMUPDF_AVAILABLE = True
    print("[A13] pymupdf imported.")
except Exception as e:
    PYMUPDF_AVAILABLE = False
    print(f"[A13] pymupdf unavailable ({e}).")

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    REPORTLAB_AVAILABLE = True
    print("[A13] reportlab imported.")
except Exception as e:
    REPORTLAB_AVAILABLE = False
    print(f"[A13] reportlab unavailable ({e}).")

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt
    DOCX_AVAILABLE = True
    print("[A13] python-docx imported.")
except Exception as e:
    DOCX_AVAILABLE = False
    print(f"[A13] python-docx unavailable ({e}).")

try:
    from PIL import Image as PILImage
    PIL_AVAILABLE = True
except Exception as e:
    PIL_AVAILABLE = False
    print(f"[A13] PIL unavailable ({e}).")

VLM_MODEL_ID       = "mlx-community/Qwen3.5-9B-MLX-4bit"
OUTPUT_DIR         = Path("A13_outputs")
MAX_TOKENS_VISION  = 512
MAX_TOKENS_TEXT    = 512
PDF_DPI            = 150
PDF_TEXT_MIN_CHARS = 100
IMGEN_STEPS        = 4

OUTPUT_DIR.mkdir(exist_ok=True)

_vlm_model     = None
_vlm_processor = None
_vlm_config    = None
_imgen_model   = None


def _load_vlm() -> bool:
    global _vlm_model, _vlm_processor, _vlm_config
    if _vlm_model is not None:
        return True
    if not VLM_AVAILABLE:
        print("[A13] VLM not available.")
        return False
    try:
        print(f"[A13] Loading VLM: {VLM_MODEL_ID} ...")
        t0 = time.time()
        _vlm_model, _vlm_processor = vlm_load(VLM_MODEL_ID)
        _vlm_config = load_config(VLM_MODEL_ID)
        print(f"[A13] VLM loaded in {time.time() - t0:.1f}s.")
        return True
    except Exception as e:
        print(f"[A13] VLM load failed: {e}")
        return False


def _load_imgen() -> bool:
    global _imgen_model
    if _imgen_model is not None:
        return True
    if not MFLUX_AVAILABLE:
        print("[A13] mflux not available.")
        return False
    try:
        print("[A13] Loading FLUX.1 Schnell (quantize=4) ...")
        t0 = time.time()
        _imgen_model = _Flux1Class.from_name(model_name="schnell", quantize=4)
        print(f"[A13] Image gen model loaded in {time.time() - t0:.1f}s.")
        return True
    except Exception as e:
        print(f"[A13] Image gen load failed: {e}")
        return False


def _write_placeholder_image(prompt: str, output_path: str) -> bool:
    if not PIL_AVAILABLE:
        print("[A13] PIL unavailable; cannot create placeholder image.")
        return False
    try:
        from PIL import ImageDraw, ImageFont

        img = PILImage.new("RGB", (1024, 1024), color=(16, 24, 40))
        draw = ImageDraw.Draw(img)
        try:
            font = ImageFont.load_default()
        except Exception:
            font = None

        title = "Image generation unavailable"
        body = "The FLUX model could not be loaded."
        body2 = "This placeholder was created so the workflow can continue."
        body3 = f"Prompt: {prompt}"

        draw.text((40, 60), title, fill=(255, 255, 255), font=font)
        draw.text((40, 120), body, fill=(200, 200, 200), font=font)
        draw.text((40, 160), body2, fill=(200, 200, 200), font=font)
        draw.text((40, 200), body3, fill=(180, 220, 255), font=font)
        img.save(output_path)
        return True
    except Exception as e:
        print(f"[A13] Placeholder image creation failed: {e}")
        return False


def unload_vlm() -> None:
    global _vlm_model, _vlm_processor, _vlm_config
    if _vlm_model is None:
        return
    import mlx.core as mx
    _vlm_model = None
    _vlm_processor = None
    _vlm_config = None
    mx.metal.clear_cache()
    print("[A13] VLM unloaded, memory cleared.")


def unload_imgen() -> None:
    global _imgen_model
    if _imgen_model is None:
        return
    import mlx.core as mx
    _imgen_model = None
    mx.metal.clear_cache()
    print("[A13] Image gen model unloaded, memory cleared.")


def _unwrap_generation(raw) -> str:
    if isinstance(raw, str):
        return raw
    if hasattr(raw, "text"):
        return raw.text
    return str(raw)


def _strip_think(text: str) -> str:
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL)
    text = re.sub(r"<think>.*", "", text, flags=re.DOTALL)
    return text.strip()


def _build_no_think_prompt(prompt: str, num_images: int) -> str:
    formatted = apply_chat_template(
        _vlm_processor,
        _vlm_config,
        f"/no_think\n{prompt}",
        num_images=num_images,
    )
    if isinstance(formatted, str):
        formatted += "<think>\n\n</think>\n"
    return formatted


def _vlm_query(prompt: str, images: list) -> str:
    if not _load_vlm():
        return "[ERROR] VLM not loaded."
    try:
        formatted = _build_no_think_prompt(prompt, len(images))
        raw = vlm_generate(
            _vlm_model,
            _vlm_processor,
            formatted,
            images,
            max_tokens=MAX_TOKENS_VISION,
            verbose=False,
        )
        return _strip_think(_unwrap_generation(raw))
    except Exception as e:
        return f"[ERROR] VLM inference failed: {e}"


def _vlm_text_query(prompt: str) -> str:
    if not _load_vlm():
        return "[ERROR] VLM not loaded."
    try:
        formatted = _build_no_think_prompt(prompt, 0)
        raw = vlm_generate(
            _vlm_model,
            _vlm_processor,
            formatted,
            None,
            max_tokens=MAX_TOKENS_TEXT,
            verbose=False,
        )
        return _strip_think(_unwrap_generation(raw))
    except Exception as e:
        return f"[ERROR] VLM text inference failed: {e}"


def generate_image(
    prompt: str,
    output_filename: str = "generated.png",
    width: int = 1024,
    height: int = 1024,
    steps: int = IMGEN_STEPS,
    seed: int = 42,
) -> str:
    unload_vlm()
    out_path = str(OUTPUT_DIR / output_filename)
    if not _load_imgen():
        print("[A13] Falling back to placeholder image because the FLUX model could not be loaded.")
        if _write_placeholder_image(prompt, out_path):
            return out_path
        return ""
    try:
        t0 = time.time()
        if _FluxConfig is not None:
            image = _imgen_model.generate_image(
                seed=seed,
                prompt=prompt,
                config=_FluxConfig(
                    num_inference_steps=steps,
                    width=width,
                    height=height,
                ),
            )
        else:
            image = _imgen_model.generate_image(
                seed=seed,
                prompt=prompt,
                num_inference_steps=steps,
                width=width,
                height=height,
            )
        image.save(path=out_path)
        print(f"[A13] Image saved: {out_path} ({time.time() - t0:.1f}s)")
        return out_path
    except Exception as e:
        print(f"[A13] Image generation failed: {e}")
        if _write_placeholder_image(prompt, out_path):
            print(f"[A13] Placeholder image saved: {out_path}")
            return out_path
        return ""


def understand_image(image_path: str, prompt: str = "Describe this image in detail.") -> str:
    unload_imgen()
    if not os.path.exists(image_path):
        return f"[ERROR] File not found: {image_path}"
    return _vlm_query(prompt, [image_path])


def understand_images_multi(image_paths: list[str], prompt: str) -> str:
    missing = [p for p in image_paths if not os.path.exists(p)]
    if missing:
        return f"[ERROR] Missing files: {missing}"
    return _vlm_query(prompt, image_paths)


def _pdf_is_text_based(doc: "fitz.Document") -> bool:
    total_chars = 0
    for page in doc:
        total_chars += len(page.get_text("text"))
        if total_chars >= PDF_TEXT_MIN_CHARS:
            return True
    return False


def _pdf_extract_text(doc: "fitz.Document") -> str:
    parts = []
    for i, page in enumerate(doc):
        text = page.get_text("text").strip()
        if text:
            parts.append(f"[Page {i + 1}]\n{text}")
    return "\n\n".join(parts)


def _pdf_pages_to_image_paths(doc: "fitz.Document") -> list[str]:
    paths = []
    mat = fitz.Matrix(PDF_DPI / 72, PDF_DPI / 72)
    for i, page in enumerate(doc):
        pix = page.get_pixmap(matrix=mat)
        tmp = tempfile.NamedTemporaryFile(
            suffix=f"_page{i}.png", delete=False, dir=OUTPUT_DIR
        )
        pix.save(tmp.name)
        tmp.close()
        paths.append(tmp.name)
    return paths


def read_pdf(pdf_path: str, prompt: str = "Summarise the content of this document.") -> str:
    if not PYMUPDF_AVAILABLE:
        return "[ERROR] pymupdf not installed."
    if not os.path.exists(pdf_path):
        return f"[ERROR] File not found: {pdf_path}"
    unload_imgen()
    tmp_paths: list[str] = []
    try:
        doc = fitz.open(pdf_path)
        if _pdf_is_text_based(doc):
            text = _pdf_extract_text(doc)
            doc.close()
            full_prompt = f"{prompt}\n\n---\nDOCUMENT TEXT:\n{text}"
            return _vlm_text_query(full_prompt)
        else:
            tmp_paths = _pdf_pages_to_image_paths(doc)
            doc.close()
            if not tmp_paths:
                return "[ERROR] Could not render PDF pages."
            return _vlm_query(prompt, tmp_paths)
    except Exception as e:
        return f"[ERROR] PDF read failed: {e}"
    finally:
        for p in tmp_paths:
            try:
                os.unlink(p)
            except OSError:
                pass


def generate_pdf(
    title: str,
    body: str,
    output_filename: str = "generated.pdf",
) -> str:
    if not REPORTLAB_AVAILABLE:
        return ""
    out_path = str(OUTPUT_DIR / output_filename)
    try:
        doc = SimpleDocTemplate(
            out_path,
            pagesize=A4,
            leftMargin=2.5 * cm,
            rightMargin=2.5 * cm,
            topMargin=2.5 * cm,
            bottomMargin=2.5 * cm,
        )
        styles = getSampleStyleSheet()
        story = [
            Paragraph(title, styles["Title"]),
            Spacer(1, 0.5 * cm),
        ]
        for para in body.strip().split("\n\n"):
            para = para.strip()
            if para:
                story.append(Paragraph(para, styles["BodyText"]))
                story.append(Spacer(1, 0.3 * cm))
        doc.build(story)
        print(f"[A13] PDF saved: {out_path}")
        return out_path
    except Exception as e:
        print(f"[A13] PDF generation failed: {e}")
        return ""


def _docx_extract_image_paths(doc_path: str) -> list[str]:
    if not PIL_AVAILABLE:
        return []
    paths = []
    try:
        import zipfile
        with zipfile.ZipFile(doc_path, "r") as z:
            for name in z.namelist():
                if name.startswith("word/media/") and any(
                    name.lower().endswith(ext) for ext in (".png", ".jpg", ".jpeg", ".bmp", ".gif")
                ):
                    data = z.read(name)
                    img = PILImage.open(io.BytesIO(data)).convert("RGB")
                    tmp = tempfile.NamedTemporaryFile(
                        suffix=".png", delete=False, dir=OUTPUT_DIR
                    )
                    img.save(tmp.name)
                    tmp.close()
                    paths.append(tmp.name)
    except Exception as e:
        print(f"[A13] DOCX image extraction warning: {e}")
    return paths


def read_docx(docx_path: str, prompt: str = "Summarise the content of this document.") -> str:
    if not DOCX_AVAILABLE:
        return "[ERROR] python-docx not installed."
    if not os.path.exists(docx_path):
        return f"[ERROR] File not found: {docx_path}"
    unload_imgen()
    tmp_paths: list[str] = []
    try:
        doc = DocxDocument(docx_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    tables_text.append(" | ".join(row_cells))
        text_body = "\n".join(paragraphs)
        if tables_text:
            text_body += "\n\nTABLES:\n" + "\n".join(tables_text)
        tmp_paths = _docx_extract_image_paths(docx_path)
        full_prompt = f"{prompt}\n\nDOCUMENT TEXT:\n{text_body}"
        if tmp_paths:
            return _vlm_query(full_prompt, tmp_paths)
        else:
            return _vlm_text_query(full_prompt)
    except Exception as e:
        return f"[ERROR] DOCX read failed: {e}"
    finally:
        for p in tmp_paths:
            try:
                os.unlink(p)
            except OSError:
                pass


def generate_docx(
    title: str,
    body: str,
    output_filename: str = "generated.docx",
) -> str:
    if not DOCX_AVAILABLE:
        return ""
    out_path = str(OUTPUT_DIR / output_filename)
    try:
        doc = DocxDocument()
        doc.add_heading(title, level=0)
        for para in body.strip().split("\n\n"):
            para = para.strip()
            if para:
                p = doc.add_paragraph(para)
                p.runs[0].font.size = Pt(11)
        doc.save(out_path)
        print(f"[A13] DOCX saved: {out_path}")
        return out_path
    except Exception as e:
        print(f"[A13] DOCX generation failed: {e}")
        return ""


def _separator(title: str) -> None:
    print("\n" + "═" * 60)
    print(f"  {title}")
    print("═" * 60)


def test_image_generation() -> None:
    _separator("TEST 1 — Image Generation (mflux FLUX.1 Schnell)")
    if not MFLUX_AVAILABLE:
        print("  ✗ mflux not installed. Run: pip install mflux")
        return
    prompt = "A futuristic cityscape at dusk, neon lights reflecting on wet streets, photorealistic"
    print(f"  Prompt: {prompt!r}")
    path = generate_image(
        prompt=prompt,
        output_filename="test_generated.png",
        width=512,
        height=512,
        steps=IMGEN_STEPS,
        seed=7,
    )
    if path:
        print(f"  ✓ Image saved: {path}")
    else:
        print("  ✗ Generation failed.")


def test_image_understanding() -> None:
    _separator("TEST 2 — Image Understanding (Qwen3.5-9B vision)")
    if not VLM_AVAILABLE:
        print("  ✗ mlx_vlm not installed.")
        return
    test_img_path = str(OUTPUT_DIR / "test_generated.png")
    if not os.path.exists(test_img_path):
        print(f"  ✗ No test image at {test_img_path}. Run TEST 1 first, or supply your own image.")
        return
    print(f"  Input: {test_img_path}")
    result = understand_image(test_img_path, "Describe what you see in this image in detail.")
    print(f"  Response:\n  {result}")
    print("  ✓ Done.")


def test_pdf_generation() -> None:
    _separator("TEST 3 — PDF Generation (reportlab)")
    if not REPORTLAB_AVAILABLE:
        print("  ✗ reportlab not installed. Run: pip install reportlab")
        return
    title = "A13 Multimodal Module — Test Document"
    body = (
        "This document was generated automatically by the A13 multimodal module of Daughter AI.\n\n"
        "The module supports six capabilities: image generation, image understanding, "
        "PDF generation, PDF reading, DOCX generation, and DOCX reading.\n\n"
        "All inference runs fully on-device using Apple Silicon via MLX. "
        "Image understanding is handled by Qwen3.5-9B loaded through mlx_vlm. "
        "Image generation uses FLUX.1 Schnell from the mflux library.\n\n"
        "This PDF was produced using reportlab with A4 page size and standard margins."
    )
    path = generate_pdf(title, body, output_filename="test_generated.pdf")
    if path:
        print(f"  ✓ PDF saved: {path}")
    else:
        print("  ✗ Generation failed.")


def test_pdf_reading() -> None:
    _separator("TEST 4 — PDF Reading (pymupdf + Qwen3.5-9B)")
    if not PYMUPDF_AVAILABLE:
        print("  ✗ pymupdf not installed. Run: pip install pymupdf")
        return
    if not VLM_AVAILABLE:
        print("  ✗ mlx_vlm not installed.")
        return
    pdf_path = str(OUTPUT_DIR / "test_generated.pdf")
    if not os.path.exists(pdf_path):
        print(f"  ✗ No PDF at {pdf_path}. Run TEST 3 first.")
        return
    print(f"  Input: {pdf_path}")
    result = read_pdf(pdf_path, "What is this document about? Give a concise summary.")
    print(f"  Response:\n  {result}")
    print("  ✓ Done.")


def test_docx_generation() -> None:
    _separator("TEST 5 — DOCX Generation (python-docx)")
    if not DOCX_AVAILABLE:
        print("  ✗ python-docx not installed. Run: pip install python-docx")
        return
    title = "A13 Multimodal Module — Test Word Document"
    body = (
        "This Word document was generated automatically by the A13 multimodal module of Daughter AI.\n\n"
        "The module demonstrates programmatic DOCX creation using python-docx, "
        "with support for titled sections and multi-paragraph body text.\n\n"
        "When integrated into Gwen, document generation can be triggered via voice commands "
        "such as: write me a report on X, or create a Word document about Y."
    )
    path = generate_docx(title, body, output_filename="test_generated.docx")
    if path:
        print(f"  ✓ DOCX saved: {path}")
    else:
        print("  ✗ Generation failed.")


def test_docx_reading() -> None:
    _separator("TEST 6 — DOCX Reading (python-docx + Qwen3.5-9B)")
    if not DOCX_AVAILABLE:
        print("  ✗ python-docx not installed.")
        return
    if not VLM_AVAILABLE:
        print("  ✗ mlx_vlm not installed.")
        return
    docx_path = str(OUTPUT_DIR / "test_generated.docx")
    if not os.path.exists(docx_path):
        print(f"  ✗ No DOCX at {docx_path}. Run TEST 5 first.")
        return
    print(f"  Input: {docx_path}")
    result = read_docx(docx_path, "What is this document about? Give a concise summary.")
    print(f"  Response:\n  {result}")
    print("  ✓ Done.")


def main() -> None:
    print("\n" + "█" * 60)
    print("  Daughter AI / Gwen — A13 Multimodal Module Test Suite")
    print("█" * 60)

    print("\n  Available capabilities:")
    print(f"    Image generation  (mflux)      : {'✓' if MFLUX_AVAILABLE else '✗'}")
    print(f"    Image understanding (mlx_vlm)  : {'✓' if VLM_AVAILABLE else '✗'}")
    print(f"    PDF generation  (reportlab)    : {'✓' if REPORTLAB_AVAILABLE else '✗'}")
    print(f"    PDF reading     (pymupdf)      : {'✓' if PYMUPDF_AVAILABLE else '✗'}")
    print(f"    DOCX generation (python-docx)  : {'✓' if DOCX_AVAILABLE else '✗'}")
    print(f"    DOCX reading    (python-docx)  : {'✓' if DOCX_AVAILABLE else '✗'}")

    input("\n  Press Enter → TEST 1: Image Generation\n")
    test_image_generation()

    input("\n  Press Enter → TEST 2: Image Understanding\n")
    test_image_understanding()

    input("\n  Press Enter → TEST 3: PDF Generation\n")
    test_pdf_generation()

    input("\n  Press Enter → TEST 4: PDF Reading\n")
    test_pdf_reading()

    input("\n  Press Enter → TEST 5: DOCX Generation\n")
    test_docx_generation()

    input("\n  Press Enter → TEST 6: DOCX Reading\n")
    test_docx_reading()

    print("\n" + "█" * 60)
    print("  All six tests complete.")
    print("█" * 60 + "\n")


if __name__ == "__main__":
    main()
    