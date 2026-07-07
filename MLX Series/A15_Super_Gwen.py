import hashlib
import io
import json
import os
import random
import re
import signal
import socket
import sqlite3
import subprocess
import tempfile
import threading
import time
from dataclasses import dataclass
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, Optional
from urllib.parse import quote

import numpy as np
import ollama
import requests
import sounddevice as sd
from mlx_lm import load as load_lm, generate as generate_lm
from audio_utils import play_audio, select_audio_device
from mlx_lm.sample_utils import make_sampler, make_logits_processors
from mlx_vlm import load as load_vlm, generate as generate_vlm

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except Exception as e:
    BS4_AVAILABLE = False
    print(f"[SuperGwen] beautifulsoup4 unavailable ({e}). Web search snippets only.")

try:
    from ddgs import DDGS
    DDGS_AVAILABLE = True
except Exception as e:
    DDGS_AVAILABLE = False
    print(f"[SuperGwen] ddgs unavailable ({e}). web_search tool will fail.")

try:
    import mlx_whisper
    WHISPER_AVAILABLE = True
except Exception as e:
    WHISPER_AVAILABLE = False
    print(f"[SuperGwen] mlx_whisper unavailable ({e}). Voice input disabled.")

try:
    from mlx_audio.tts.models.kokoro import KokoroPipeline
    from mlx_audio.tts.utils import load_model as load_tts_model
    TTS_AVAILABLE = True
except Exception as e:
    TTS_AVAILABLE = False
    print(f"[SuperGwen] mlx_audio unavailable ({e}). TTS disabled, text-only replies.")

try:
    import chromadb
    from chromadb.utils import embedding_functions
    CHROMADB_IMPORTABLE = True
except Exception as e:
    CHROMADB_IMPORTABLE = False
    print(f"[SuperGwen] chromadb unavailable ({e}). Semantic memory disabled.")

try:
    import fitz  # pymupdf
    PYMUPDF_AVAILABLE = True
except Exception as e:
    PYMUPDF_AVAILABLE = False
    print(f"[SuperGwen] pymupdf unavailable ({e}). PDF reading disabled.")

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    REPORTLAB_AVAILABLE = True
except Exception as e:
    REPORTLAB_AVAILABLE = False
    print(f"[SuperGwen] reportlab unavailable ({e}). PDF generation disabled.")

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except Exception as e:
    DOCX_AVAILABLE = False
    print(f"[SuperGwen] python-docx unavailable ({e}). DOCX read/write disabled.")

try:
    from PIL import Image as PILImage, ImageDraw, ImageFont
    PIL_AVAILABLE = True
except Exception as e:
    PIL_AVAILABLE = False
    print(f"[SuperGwen] Pillow unavailable ({e}).")

try:
    from mflux.models.flux.variants.txt2img.flux import Flux1 as _Flux1Class
    MFLUX_AVAILABLE = True
except Exception as e:
    _Flux1Class = None
    MFLUX_AVAILABLE = False
    print(f"[SuperGwen] mflux unavailable ({e}). Image generation will use a placeholder.")

try:
    from A14_VoiceGate import VoiceGate
    VOICEGATE_IMPORTABLE = True
except Exception as e:
    VoiceGate = None
    VOICEGATE_IMPORTABLE = False
    print(f"[SuperGwen] A14_VoiceGate unavailable ({e}). Voice-gate auth disabled.")

try:
    import soundfile as sf
    SOUNDFILE_AVAILABLE = True
except Exception as e:
    SOUNDFILE_AVAILABLE = False


# ------
# SIGNAL HANDLING
# ------

def _signal_handler(signum, frame):
    """Handle signals gracefully."""
    try:
        sd.stop()
    except Exception:
        pass
    print(f"\n[Signal] Received signal {signum}, shutting down...")
    exit(0)

signal.signal(signal.SIGTRAP, _signal_handler)
signal.signal(signal.SIGINT, _signal_handler)
signal.signal(signal.SIGTERM, _signal_handler)


# -------
# CONFIG
# -------

DB_PATH = os.path.expanduser("~/daughter_ai.db")           
CHROMA_PATH = os.path.expanduser("~/daughter_ai_chroma")
OUTPUT_DIR = Path("SuperGwen_outputs")
OUTPUT_DIR.mkdir(exist_ok=True)

ENABLE_VOICE_GATE = True
VOICEPRINT_PATH = "Gwen_voiceprint.pt"
VOICE_GATE_THRESHOLD = 0.30
GUARDED_TOOLS = {"quit_app", "update_memory"}

DEFAULT_MODEL_TIER_NAME = "REGULAR"
DEFAULT_INPUT_MODE = "voice"

SAMPLE_RATE = 16_000
CHUNK_SIZE = 1_024
SILENCE_THRESHOLD = 0.01
SILENCE_DURATION = 1.5
WAKE_TIMEOUT = 30.0
CACHE_TTL_HOURS = 6
VLM_MAX_TOKENS = 512


# --------
# LOGGING
# --------

def _ts() -> str:
    return datetime.now().strftime("%H:%M:%S")

def log_success(msg: str) -> None: print(f"[SUCCESS {_ts()}] {msg}")
def log_info(msg: str) -> None:    print(f"[INFO {_ts()}] {msg}")
def log_error(msg: str) -> None:   print(f"[ERROR {_ts()}] {msg}")
def log_failed(msg: str) -> None:  print(f"[FAILED {_ts()}] {msg}")
def log_skipped(msg: str) -> None: print(f"[SKIPPED {_ts()}] {msg}")


# ------------------------------------
# MODEL TIERS ( MANUAL SELECTION ONLY)
# ------------------------------------

class ModelTier(Enum):
    VERY_LIGHT = "Very Light"
    LIGHT = "Light"
    REGULAR = "Regular"
    MAX = "MAX"


@dataclass(frozen=True)
class ModelConfig:
    name: str
    backend: str          # "mlx_lm", "mlx_vlm", "ollama"
    family: str           # "qwen", "gemma", "kimi"
    temperature: float = 0.0
    think: bool = False    # only for ollama backend


MODEL_REGISTRY: dict[ModelTier, ModelConfig] = {
    ModelTier.VERY_LIGHT: ModelConfig(
        name="mlx-community/Qwen2.5-1.5B-Instruct-4bit",
        backend="mlx_lm",
        family="qwen",
        temperature=0.0,
    ),
    ModelTier.LIGHT: ModelConfig(
        name="mlx-community/Qwen3.5-9B-MLX-4bit",
        backend="mlx_vlm",
        family="qwen",
        temperature=0.0,
    ),
    ModelTier.REGULAR: ModelConfig(
        name="gemma4:cloud",
        backend="ollama",
        family="gemma",
        temperature=0.1,
        think=False,
    ),
    ModelTier.MAX: ModelConfig(
        name="kimi-k2.7-code:cloud",
        backend="ollama",
        family="kimi",
        temperature=0.1,
        think=False,
    ),
}

REASONING_PATTERNS: dict[str, list[str]] = {
    "qwen": [r"<think>.*?</think>"],
    "gemma": [r"<think>.*?</think>", r"<thinking>.*?</thinking>", r"<reasoning>.*?</reasoning>"],
    "kimi": [r"<think>.*?</think>", r"<thinking>.*?</thinking>", r"<reasoning>.*?</reasoning>"],
}


# ------------------------------
# NETWORK / CLOUD AVAILABILITY
# ------------------------------

def ollama_available() -> bool:
    try:
        ollama.list()
        return True
    except Exception:
        return False

def network_connection_available(timeout: float = 3.0) -> bool:
    try:
        socket.create_connection(("8.8.8.8", 53), timeout=timeout)
        return True
    except OSError:
        return False

def can_use_ollama_cloud() -> bool:
    return network_connection_available() and ollama_available()


# ----------------------
# TEXT CLEANING HELPERS
# ----------------------

def clean_response(text: str, family: str) -> str:
    for pattern in REASONING_PATTERNS.get(family, []):
        text = re.sub(pattern, "", text, flags=re.DOTALL | re.IGNORECASE)
    return text.strip()

def strip_special_tokens(text: str) -> str:
    text = re.sub(r"<\|im_start\|>\s*assistant\s*", "", text)
    text = re.sub(r"<\|im_end\|>", "", text)
    return text.strip()

def extract_mlx_text(response: Any) -> str:
    return response.text if hasattr(response, "text") else str(response)

def extract_ollama_text(response: Any) -> str:
    if hasattr(response, "message"):
        message = response.message
        if hasattr(message, "content"):
            return message.content or ""
    if isinstance(response, dict):
        message = response.get("message", {})
        if isinstance(message, dict):
            return message.get("content", "") or ""
    return str(response)

def sanitize_response(text: str) -> str:
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL)
    text = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", text)
    text = re.sub(r"#{1,6}\s*", "", text)
    text = re.sub(r"`{1,3}[^`]*`{1,3}", "", text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    text = re.sub(r"[_~>|*]", " ", text)
    text = text.encode("ascii", errors="ignore").decode("ascii")
    text = re.sub(r"\s+", " ", text).strip()
    return text


# -------------
# MODEL ROUTER 
# -------------

class ModelRouter:
    def __init__(self):
        self.loaded_models: dict[str, tuple[Any, Any]] = {}
        self.current_tier: Optional[ModelTier] = None

    def _load_backend(self, tier: ModelTier) -> bool:
        config = MODEL_REGISTRY[tier]

        if config.backend == "ollama":
            if not can_use_ollama_cloud():
                log_error(f"Cannot reach Ollama cloud for {config.name}. Check network / `ollama signin`.")
                return False
            return True

        if config.name in self.loaded_models:
            log_info(f"Using cached model: {config.name}")
            return True

        try:
            if config.backend == "mlx_lm":
                model, tokenizer = load_lm(config.name)
                self.loaded_models[config.name] = (model, tokenizer)
            elif config.backend == "mlx_vlm":
                model, processor = load_vlm(config.name)
                self.loaded_models[config.name] = (model, processor)
            else:
                raise ValueError(f"Unknown backend: {config.backend}")
            log_success(f"Loaded model: {config.name}")
            return True
        except Exception as e:
            log_error(f"Failed to load {config.name}: {e}")
            return False

    def load(self, tier: ModelTier) -> bool:
        ok = self._load_backend(tier)
        if ok:
            self.current_tier = tier
        return ok

    def ensure_loaded(self, tier: ModelTier) -> bool:
        return self._load_backend(tier)

    def unload(self, tier: ModelTier) -> None:
        config = MODEL_REGISTRY[tier]
        if config.name in self.loaded_models:
            del self.loaded_models[config.name]
            try:
                import mlx.core as mx
                mx.metal.clear_cache()
            except Exception:
                pass
            log_info(f"Unloaded {config.name}, memory cleared.")

    def get_cached(self, tier: ModelTier) -> tuple[Any, Any]:
        config = MODEL_REGISTRY[tier]
        return self.loaded_models[config.name]

    # chat dispatch

    def chat(self, system_prompt: str, history: list[dict], user_input: str) -> str:
        if self.current_tier is None:
            raise RuntimeError("No model tier selected. Call load() first.")
        config = MODEL_REGISTRY[self.current_tier]

        if config.backend == "ollama":
            try:
                messages = self._build_ollama_messages(system_prompt, history, user_input)
                return self._generate_ollama_chat(messages, config)
            
            except RuntimeError as e:
                fallback_cfg = MODEL_REGISTRY[ModelTier.LIGHT]
                log_error(f"{config.name} unavailable ({e}). Falling back to {fallback_cfg.name}.")
                if self.load(ModelTier.LIGHT):
                    prompt = self._build_mlx_prompt(system_prompt, history, user_input)
                    return self._generate_mlx_vlm(prompt, fallback_cfg)
                raise

        if config.backend == "mlx_lm":
            prompt = self._build_mlx_prompt(system_prompt, history, user_input)
            return self._generate_mlx_lm(prompt, config)
        
        if config.backend == "mlx_vlm":
            prompt = self._build_mlx_prompt(system_prompt, history, user_input)
            return self._generate_mlx_vlm(prompt, config)
        
        raise ValueError(f"Unknown backend: {config.backend}")

    @staticmethod
    def _build_mlx_prompt(system_prompt: str, history: list[dict], user_input: str) -> str:
        prompt = f"<|im_start|>system\n{system_prompt}<|im_end|>\n"
        for m in history:
            prompt += f"<|im_start|>{m['role']}\n{m['content']}<|im_end|>\n"
        prompt += f"<|im_start|>user\n{user_input}<|im_end|>\n<|im_start|>assistant\n"
        return prompt

    @staticmethod
    def _build_ollama_messages(system_prompt: str, history: list[dict], user_input: str) -> list[dict]:
        messages = [{"role": "system", "content": system_prompt}]
        messages.extend(history)
        messages.append({"role": "user", "content": user_input})
        return messages

    def _generate_mlx_lm(self, prompt: str, config: ModelConfig) -> str:
        model, tokenizer = self.loaded_models[config.name]
        response = generate_lm(
            model=model,
            tokenizer=tokenizer,
            prompt=prompt,
            sampler=make_sampler(temp=config.temperature),
            verbose=False,
        )
        raw = extract_mlx_text(response)
        return strip_special_tokens(clean_response(raw, config.family))

    def _generate_mlx_vlm(self, prompt: str, config: ModelConfig) -> str:
        model, processor = self.loaded_models[config.name]
        response = generate_vlm(
            model=model,
            processor=processor,
            prompt=prompt,
            temperature=config.temperature,
            verbose=False,
        )
        raw = extract_mlx_text(response)
        return strip_special_tokens(clean_response(raw, config.family))

    def _generate_ollama_chat(self, messages: list[dict], config: ModelConfig) -> str:
        try:
            response = ollama.chat(
                model=config.name,
                messages=messages,
                think=config.think,
                options={"temperature": config.temperature},
            )
            raw = extract_ollama_text(response)
            return clean_response(raw, config.family)
        except ollama.ResponseError as e:
            if e.status_code == 401:
                raise RuntimeError("Ollama auth failed. Run: ollama signin") from e
            if e.status_code == 429:
                raise RuntimeError("Ollama cloud limit hit. Try a local tier for now.") from e
            raise RuntimeError(f"Ollama error {e.status_code}: {e}") from e


router = ModelRouter()


# ----------------------------
# MANUAL MODEL-TIER SELECTION
# ----------------------------

class ModelModeController:
    PHRASE_TO_TIER = {
        "very light": ModelTier.VERY_LIGHT,
        "very light mode": ModelTier.VERY_LIGHT,
        "light": ModelTier.LIGHT,
        "light mode": ModelTier.LIGHT,
        "regular": ModelTier.REGULAR,
        "regular mode": ModelTier.REGULAR,
        "default mode": ModelTier.REGULAR,
        "max": ModelTier.MAX,
        "max mode": ModelTier.MAX,
        "high mode": ModelTier.MAX,
    }

    STATUS_PHRASES = {"what model", "which model", "what mode are you in", "current model", "which mode"}

    def __init__(self, router: ModelRouter, initial: ModelTier):
        self.router = router
        if router.load(initial):
            return
        
        log_error(f"Could not load startup tier {initial.value}; falling back to {ModelTier.LIGHT.value}")
        
        if not router.load(ModelTier.LIGHT):
            raise RuntimeError(f"Could not load fallback tier {ModelTier.LIGHT.value} either.")

    @property
    def tier(self) -> ModelTier:
        return self.router.current_tier

    def _normalize(self, utterance: str) -> Optional[str]:
        lowered = utterance.strip().lower()
        if lowered.startswith("/model"):
            lowered = lowered.replace("/model", "").strip()
        elif lowered.startswith("switch to "):
            lowered = lowered[len("switch to "):].strip()
        elif lowered.startswith("use "):
            lowered = lowered[len("use "):].strip()
        return lowered

    def is_status_query(self, utterance: str) -> bool:
        return utterance.strip().lower() in self.STATUS_PHRASES

    def status_text(self) -> str:
        return f"I'm currently running in {self.tier.value} mode."

    def process(self, utterance: str) -> Optional[str]:
        lowered = self._normalize(utterance)
        if lowered not in self.PHRASE_TO_TIER:
            return None

        requested = self.PHRASE_TO_TIER[lowered]
        if requested == self.tier:
            return f"Already in {requested.value} mode."

        if self.router.load(requested):
            log_success(f"Model tier switched to {requested.value}")
            return f"Switched to {requested.value} mode."
        else:
            return (
                f"Couldn't switch to {requested.value} mode — cloud model unreachable. "
                f"Staying on {self.tier.value} mode."
            )


# ------------------------
# TEXT / VOICE INPUT MODE
# ------------------------

class InputModeController:
    VOICE = "voice"
    TEXT = "text"

    _TO_TEXT = {"text mode", "switch to text", "switch to text mode", "type mode"}
    _TO_VOICE = {"voice mode", "back to voice", "voice input"}

    def __init__(self, initial: str = VOICE):
        self.mode = initial
        self._print_mode()

    def _print_mode(self):
        icon = "\U0001F399" if self.mode == self.VOICE else "\u2328"
        print(f"  [{icon} INPUT: {self.mode.upper()}]")

    def process(self, utterance: str) -> bool:
        lowered = utterance.strip().lower()
        if lowered in self._TO_TEXT and self.mode != self.TEXT:
            self.mode = self.TEXT
            self._print_mode()
            return True
        if lowered in self._TO_VOICE and self.mode != self.VOICE:
            self.mode = self.VOICE
            self._print_mode()
            return True
        return False

    @property
    def is_voice(self) -> bool:
        return self.mode == self.VOICE


# ------
# AUDIO
# ------

if TTS_AVAILABLE:
    _voice_model = load_tts_model("prince-canuma/Kokoro-82M")
    _tts_pipeline = KokoroPipeline(lang_code="a", model=_voice_model, repo_id="prince-canuma/Kokoro-82M")
else:
    _tts_pipeline = None

TTS_RATE = 24_000
INTERRUPT_RMS_THRESHOLD = SILENCE_THRESHOLD * 3.0
MIC_BLOCK_SIZE = 512
_interrupt_event = threading.Event()

ACTIVATION_PHRASES = [
    "I'm listening.", "Hey.", "Yeah?", "Uh huh.", "Go ahead.",
    "How can I help?", "What's up?", "Tell me.",
]

def pick_activation_phrase() -> str:
    return random.choice(ACTIVATION_PHRASES)


def _bell_tone(frequencies: list[float], duration: float = 0.18,
               sample_rate: int = 44_100, amplitude: float = 0.25) -> np.ndarray:
    n = int(sample_rate * duration)
    t = np.linspace(0, duration, n, endpoint=False)
    attack = int(0.002 * sample_rate)
    envelope = np.exp(-t / (duration * 0.45))
    envelope[:attack] = np.linspace(0, 1, attack)
    wave = np.zeros(n)
    for freq in frequencies:
        for h, w in enumerate([1.0, 0.35, 0.15, 0.06], start=1):
            wave += w * np.sin(2 * np.pi * freq * h * t)
    wave *= envelope
    peak = np.max(np.abs(wave))
    if peak > 0:
        wave = wave / peak * amplitude
    return np.stack([wave, wave], axis=1).astype(np.float32)

def beep_activate(sample_rate: int = 44_100) -> None:
    try:
        gap = np.zeros((int(sample_rate * 0.04), 2), dtype=np.float32)
        clip = np.concatenate([_bell_tone([880.0], 0.16, sample_rate), gap, _bell_tone([1320.0], 0.20, sample_rate)])
        play_audio(sd, clip, sample_rate)
    except Exception as e:
        log_error(f"beep_activate failed: {e}")

def beep_deactivate(sample_rate: int = 44_100) -> None:
    try:
        play_audio(sd, _bell_tone([523.25], 0.22, sample_rate), sample_rate)
    except Exception as e:
        log_error(f"beep_deactivate failed: {e}")

def beep_timeout(sample_rate: int = 44_100) -> None:
    try:
        gap = np.zeros((int(sample_rate * 0.05), 2), dtype=np.float32)
        clip = np.concatenate([_bell_tone([1047.0], 0.16, sample_rate), gap, _bell_tone([659.0], 0.20, sample_rate)])
        play_audio(sd, clip, sample_rate)
    except Exception as e:
        log_error(f"beep_timeout failed: {e}")


def resample_audio(audio: np.ndarray, orig_sr: int, target_sr: int) -> np.ndarray:
    if orig_sr == target_sr or audio.size == 0:
        return audio
    ratio = target_sr / orig_sr
    n = int(round(audio.shape[0] * ratio))
    idx = np.arange(n) / ratio
    left = np.floor(idx).astype(np.int64)
    right = np.minimum(left + 1, audio.shape[0] - 1)
    frac = idx - left
    return ((1 - frac) * audio[left] + frac * audio[right]).astype(audio.dtype)


def _mic_rms_monitor(stop_event: threading.Event) -> None:
    def _callback(indata, frames, time_info, status):
        if stop_event.is_set():
            raise sd.CallbackStop()
        rms = float(np.sqrt(np.mean(indata ** 2)))
        if rms > INTERRUPT_RMS_THRESHOLD:
            _interrupt_event.set()
            raise sd.CallbackStop()
    try:
        with sd.InputStream(samplerate=SAMPLE_RATE, channels=1, dtype="float32",
                     blocksize=MIC_BLOCK_SIZE, device=0, callback=_callback):
            while not stop_event.is_set() and not _interrupt_event.is_set():
                time.sleep(0.01)
    except (sd.CallbackStop, Exception):
        pass
    finally:
        try:
            sd.stop()
        except Exception:
            pass


def speak_interruptible(text: str) -> bool:
    if not TTS_AVAILABLE or _tts_pipeline is None:
        print(f"  [TTS unavailable] {text}")
        return False

    _interrupt_event.clear()
    stop_monitor = threading.Event()
    monitor = threading.Thread(target=_mic_rms_monitor, args=(stop_monitor,), daemon=True)
    monitor.start()

    sd.stop()
    time.sleep(0.05)
    
    # Pre-select a valid output device to avoid PortAudio errors
    safe_device = select_audio_device(sd, kind="output")
    
    try:
        for sentence in re.split(r"(?<=[.!?])\s+", text):
            sentence = sentence.strip()
            if not sentence:
                continue
            for _, _, audio in _tts_pipeline(sentence, voice="af_heart", speed=1.3):
                samples = np.array(audio[0], dtype=np.float32)
                try:
                    dev_info = sd.query_devices(safe_device, "output") if safe_device is not None else {}
                    dev_rate = int(dev_info.get("default_samplerate", TTS_RATE))
                except Exception:
                    dev_rate = TTS_RATE
                if dev_rate != TTS_RATE:
                    samples = resample_audio(samples, TTS_RATE, dev_rate)
                try:
                    play_audio(sd, samples, dev_rate, device_index=safe_device)
                except Exception as e:
                    log_error(f"play_audio failed: {e}")
                if _interrupt_event.is_set():
                    return True
    except Exception as e:
        log_error(f"speak_interruptible: {e}")
        sd.stop()
    finally:
        stop_monitor.set()
        monitor.join(timeout=1.0)
        sd.stop()
        interrupted = _interrupt_event.is_set()
        _interrupt_event.clear()
    return interrupted


# ----
# STT
# ----

WAKE_WORDS = {"gwen", "hey gwen", "hi gwen", "okay gwen"}
DEACTIVATION_PHRASES = {"deactivate", "go to sleep", "stop listening", "sleep", "go to sleep gwen"}
EXIT_PHRASES = {"exit", "quit", "shut down", "shutdown"}
WHISPER_REPO = "mlx-community/distil-whisper-large-v3"


def transcribe_audio(audio_data: np.ndarray) -> str:
    if not WHISPER_AVAILABLE or audio_data.size == 0:
        return ""
    result = mlx_whisper.transcribe(
        audio_data, path_or_hf_repo=WHISPER_REPO, language="en",
        task="transcribe", word_timestamps=True, verbose=False,
    )
    return result.get("text", "")


def is_wake_word(text: str) -> bool:
    t = text.strip().lower()
    if any(w in t for w in WAKE_WORDS):
        return True
    return bool(re.search(r"\bgwen\b", t))

def is_deactivation(text: str) -> bool:
    return any(p in text.strip().lower() for p in DEACTIVATION_PHRASES)

def is_exit(text: str) -> bool:
    return any(p in text.strip().lower() for p in EXIT_PHRASES)

def listen_for_wake_word() -> None:
    print("[Gwen] Dormant — say 'Hey Gwen' to activate.")
    with sd.InputStream(samplerate=SAMPLE_RATE, dtype="float32", channels=1,
                         blocksize=CHUNK_SIZE, device=0) as stream:
        while True:
            data, overflowed = stream.read(CHUNK_SIZE)
            if np.sqrt(np.mean(data ** 2)) <= SILENCE_THRESHOLD:
                continue
            frames = [data.copy()]
            silent_chunks, total = 0, 1
            max_silent = int(0.7 * SAMPLE_RATE / CHUNK_SIZE)
            max_chunks = int(2.5 * SAMPLE_RATE / CHUNK_SIZE)
            while total < max_chunks:
                chunk, _ = stream.read(CHUNK_SIZE)
                frames.append(chunk.copy())
                total += 1
                if np.sqrt(np.mean(chunk ** 2)) < SILENCE_THRESHOLD:
                    silent_chunks += 1
                    if silent_chunks >= max_silent:
                        break
                else:
                    silent_chunks = 0
            text = transcribe_audio(np.concatenate(frames).flatten())
            if text.strip():
                print(f"[Wake] Heard: {text.strip()}")
                if is_wake_word(text):
                    return


def listen_for_prompt() -> tuple[Optional[str], float]:
    deadline = time.time() + WAKE_TIMEOUT
    print(f"[Gwen] Active — waiting for your prompt ({int(WAKE_TIMEOUT)}s timeout)...")
    time.sleep(0.25)
    with sd.InputStream(samplerate=SAMPLE_RATE, dtype="float32", channels=1,
                         blocksize=CHUNK_SIZE, device=0) as stream:
        while time.time() < deadline:
            data, _ = stream.read(CHUNK_SIZE)
            if np.sqrt(np.mean(data ** 2)) <= SILENCE_THRESHOLD:
                continue
            frames = [data.copy()]
            silent_chunks = 0
            max_silent = int(SILENCE_DURATION * SAMPLE_RATE / CHUNK_SIZE)
            while True:
                chunk, _ = stream.read(CHUNK_SIZE)
                frames.append(chunk.copy())
                chunk_rms = np.sqrt(np.mean(chunk ** 2))
                peak = np.max(np.abs(chunk))
                if chunk_rms < SILENCE_THRESHOLD and peak < SILENCE_THRESHOLD * 3:
                    silent_chunks += 1
                    if silent_chunks >= max_silent:
                        break
                else:
                    silent_chunks = 0
            t_stt = time.time()
            transcription = transcribe_audio(np.concatenate(frames).flatten())
            return transcription, (time.time() - t_stt) * 1000
    return None, 0.0


# ---------
# DATABASE
# ---------

def get_db() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db()
    c = conn.cursor()
    c.execute("""CREATE TABLE IF NOT EXISTS user_facts (
        key TEXT PRIMARY KEY, value TEXT NOT NULL, updated_at TEXT NOT NULL)""")
    c.execute("""CREATE TABLE IF NOT EXISTS user_facts_history (
        id INTEGER PRIMARY KEY AUTOINCREMENT, key TEXT NOT NULL,
        old_value TEXT, new_value TEXT NOT NULL, changed_at TEXT NOT NULL)""")
    c.execute("""CREATE TABLE IF NOT EXISTS search_cache (
        query_hash TEXT PRIMARY KEY, query TEXT NOT NULL,
        answer TEXT NOT NULL, cached_at TEXT NOT NULL)""")
    c.execute("""CREATE TABLE IF NOT EXISTS session_summaries (
        id INTEGER PRIMARY KEY AUTOINCREMENT, summary TEXT NOT NULL, created_at TEXT NOT NULL)""")
    c.execute("""CREATE TABLE IF NOT EXISTS latency_log (
        id INTEGER PRIMARY KEY AUTOINCREMENT, session_id TEXT NOT NULL,
        query_preview TEXT NOT NULL, complexity TEXT NOT NULL DEFAULT 'unknown',
        model_tier TEXT NOT NULL DEFAULT 'unknown',
        stt_ms REAL, memory_ms REAL, inference_ms REAL, tts_ms REAL, total_ms REAL,
        logged_at TEXT NOT NULL)""")
    c.execute("""CREATE TABLE IF NOT EXISTS memory_retrieval_log (
        id INTEGER PRIMARY KEY AUTOINCREMENT, session_id TEXT NOT NULL,
        query_preview TEXT NOT NULL, facts_injected TEXT, semantic_hits TEXT,
        logged_at TEXT NOT NULL)""")

    existing_cols = {row["name"] for row in c.execute("PRAGMA table_info(latency_log)").fetchall()}
    if "complexity" not in existing_cols:
        c.execute("ALTER TABLE latency_log ADD COLUMN complexity TEXT NOT NULL DEFAULT 'unknown'")
        log_info("Migrated latency_log: added complexity column.")
    if "model_tier" not in existing_cols:
        c.execute("ALTER TABLE latency_log ADD COLUMN model_tier TEXT NOT NULL DEFAULT 'unknown'")
        log_info("Migrated latency_log: added model_tier column.")

    conn.commit()
    conn.close()

init_db()

SESSION_ID = datetime.now().strftime("%Y%m%d_%H%M%S")


# ChromaDB semantic memory

CHROMA_AVAILABLE = False
chroma_collection = None
if CHROMADB_IMPORTABLE:
    try:
        chroma_client = chromadb.PersistentClient(path=CHROMA_PATH)
        ef = embedding_functions.DefaultEmbeddingFunction()
        chroma_collection = chroma_client.get_or_create_collection(name="daughter_memory", embedding_function=ef)
        CHROMA_AVAILABLE = True
        log_success("ChromaDB loaded.")
    except Exception as e:
        log_error(f"ChromaDB init failed: {e}")

def store_semantic_memory(text: str, metadata: Optional[dict] = None):
    if not CHROMA_AVAILABLE:
        return
    doc_id = hashlib.md5((text + datetime.now().isoformat()).encode()).hexdigest()
    chroma_collection.add(documents=[text], metadatas=[metadata or {}], ids=[doc_id])

def retrieve_semantic_memory(query: str, n_results: int = 3) -> list[str]:
    if not CHROMA_AVAILABLE:
        return []
    try:
        count = chroma_collection.count()
        if count == 0:
            return []
        results = chroma_collection.query(query_texts=[query], n_results=min(n_results, count))
        return results["documents"][0] if results["documents"] else []
    except Exception as e:
        log_error(f"Semantic retrieval failed: {e}")
        return []


# user facts

def get_all_facts() -> dict:
    conn = get_db()
    rows = conn.execute("SELECT key, value FROM user_facts").fetchall()
    conn.close()
    return {r["key"]: r["value"] for r in rows}

def set_fact(key: str, value: str):
    conn = get_db()
    now = datetime.now().isoformat()
    existing = conn.execute("SELECT value FROM user_facts WHERE key = ?", (key,)).fetchone()
    old_value = existing["value"] if existing else None
    conn.execute("INSERT OR REPLACE INTO user_facts (key, value, updated_at) VALUES (?, ?, ?)", (key, value, now))
    conn.execute("INSERT INTO user_facts_history (key, old_value, new_value, changed_at) VALUES (?, ?, ?, ?)",
                 (key, old_value, value, now))
    conn.commit()
    conn.close()

def build_global_context() -> str:
    facts = get_all_facts()
    if not facts:
        return ""
    lines = ["[User context]"] + [f"  {k}: {v}" for k, v in facts.items()]
    return "\n".join(lines)


# auto fact extraction

SMALL_MODEL_REPO = MODEL_REGISTRY[ModelTier.VERY_LIGHT].name
try:
    small_model, small_tokenizer = load_lm(SMALL_MODEL_REPO)
    SMALL_MODEL_AVAILABLE = True
    router.loaded_models[SMALL_MODEL_REPO] = (small_model, small_tokenizer)
except Exception as e:
    log_error(f"Small model unavailable ({e}). Fact extraction disabled.")
    SMALL_MODEL_AVAILABLE = False

AUTO_UPDATE_PROMPT = """You are a fact extractor. Given a user message, extract any personal facts the user is stating about themselves (name, location, current projects, stopped projects, preferences, occupation, etc.).

Return a JSON array of objects with "key" and "value". Use snake_case keys. If no facts are stated, return [].
Return ONLY the JSON array, nothing else.

Examples:
- "I've stopped working on project X" -> [{{"key": "stopped_projects", "value": "project X"}}]
- "I moved to Chandigarh" -> [{{"key": "location", "value": "Chandigarh"}}]
- "my name is Veer" -> [{{"key": "name", "value": "Veer"}}]
- "hey what's up" -> []

User message: {message}"""


def detect_and_apply_fact_updates(user_message: str):
    if not SMALL_MODEL_AVAILABLE:
        return
    prompt = AUTO_UPDATE_PROMPT.format(message=user_message)
    try:
        response = generate_lm(
            small_model, small_tokenizer,
            prompt=f"<|im_start|>user\n{prompt}<|im_end|>\n<|im_start|>assistant\n",
            max_tokens=200, sampler=make_sampler(temp=0.0), verbose=False,
        )
        raw = clean_response(extract_mlx_text(response), "qwen")
        log_info(f"[FactExtract] raw model output: {raw!r}")   # <-- add this
        start, end = raw.find("["), raw.rfind("]")
        if start == -1 or end == -1:
            log_error(f"[FactExtract] No JSON array found in output.")   # <-- add this
            return
        for f in json.loads(raw[start:end + 1]):
            if "key" in f and "value" in f:
                set_fact(f["key"], f["value"])
                store_semantic_memory(f"User fact — {f['key']}: {f['value']}", {"type": "user_fact", "key": f["key"]})
                log_info(f"Updated fact: {f['key']} = {f['value']}")
    except Exception as e:
        log_error(f"Auto-update failed: {e}")


# -----------
# WEB SEARCH
# -----------

FORCE_SEARCH_PATTERNS = [
    r"\b(world cup|premier league|champions league|nba|nfl)\b",
    r"\b(score|standings|match|tournament|who (won|is winning|is leading))\b",
    r"\b(price|stock|weather|news|latest|current|right now|today)\b",
    r"\b(born on|birthday|who is|how old is|when did|what year)\b",
    r"\b(travel time|distance|travel|cost)\b",
]

def _query_hash(query: str) -> str:
    return hashlib.md5(query.strip().lower().encode()).hexdigest()

def get_cached_answer(query: str) -> Optional[str]:
    conn = get_db()
    row = conn.execute("SELECT answer, cached_at FROM search_cache WHERE query_hash = ?", (_query_hash(query),)).fetchone()
    conn.close()
    if not row:
        return None
    age_hours = (datetime.now() - datetime.fromisoformat(row["cached_at"])).total_seconds() / 3600
    return None if age_hours > CACHE_TTL_HOURS else row["answer"]

def cache_answer(query: str, answer: str):
    conn = get_db()
    conn.execute("INSERT OR REPLACE INTO search_cache (query_hash, query, answer, cached_at) VALUES (?, ?, ?, ?)",
                 (_query_hash(query), query, answer, datetime.now().isoformat()))
    conn.commit()
    conn.close()

deep_search_mode = False

def fetch_page_text(url: str, max_chars: int = 1500) -> str:
    if not BS4_AVAILABLE:
        return ""
    try:
        r = requests.get(url, timeout=5, headers={"User-Agent": "Mozilla/5.0"})
        soup = BeautifulSoup(r.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        return soup.get_text(separator=" ", strip=True)[:max_chars]
    except Exception:
        return ""

def get_response_with_context(user_input: str, context: str) -> str:
    augmented = f"{user_input}\n\n[Web context]:\n{context}"
    system = ("You are a helpful assistant. Answer using ONLY the provided web context. "
              "Be concise and accurate. If the context doesn't answer the question, say so.")
    raw = router.chat(system_prompt=system, history=[], user_input=augmented)
    return sanitize_response(raw)

def do_web_search(query: str, user_input: str, force_fresh: bool = False) -> str:
    if not force_fresh:
        cached = get_cached_answer(query)
        if cached:
            log_info(f"Cache hit for: {query}")
            return cached
    if not DDGS_AVAILABLE:
        return "Web search isn't available right now — the ddgs package isn't installed."

    log_info(f"Fetching: {query}")
    variants = [query, f"{query} {datetime.now().strftime('%Y')}",
                f"{query} today {datetime.now().strftime('%B %d %Y')}"]
    all_context, seen_urls = [], set()
    for q in variants:
        with DDGS() as ddgs:
            try:
                for r in list(ddgs.text(q, max_results=2)):
                    url = r.get("href", "")
                    if url in seen_urls:
                        continue
                    seen_urls.add(url)
                    snippet = r.get("body", "")
                    page_text = fetch_page_text(url) if deep_search_mode else ""
                    all_context.append(f"Source: {url}\nSnippet: {snippet}\n{page_text}")
            except Exception:
                continue
    context = "\n\n---\n\n".join(all_context) if all_context else "No results found."
    answer = sanitize_response(get_response_with_context(user_input, context))
    cache_answer(query, answer)
    store_semantic_memory(f"Web search: {query} -> {answer[:200]}", {"type": "web_search", "query": query})
    return answer


# ----------------------------------------
# IMAGE / PDF / DOCX GENERATION + READING
# ----------------------------------------

def _get_vision_model_processor():
    if not router.ensure_loaded(ModelTier.LIGHT):
        return None, None
    return router.get_cached(ModelTier.LIGHT)

def _vlm_raw(prompt: str, images: Optional[list]) -> str:
    model, processor = _get_vision_model_processor()
    if model is None:
        return "[ERROR] Vision model unavailable."
    try:
        formatted = f"<|im_start|>user\n{prompt}<|im_end|>\n<|im_start|>assistant\n"
        raw = generate_vlm(model, processor, formatted, images, max_tokens=VLM_MAX_TOKENS, verbose=False)
        return strip_special_tokens(clean_response(extract_mlx_text(raw), "qwen"))
    except Exception as e:
        return f"[ERROR] Vision inference failed: {e}"


_imgen_model = None

def _load_imgen() -> bool:
    global _imgen_model
    if _imgen_model is not None:
        return True
    if not MFLUX_AVAILABLE:
        return False
    try:
        _imgen_model = _Flux1Class.from_name(model_name="schnell", quantize=4)
        return True
    except Exception as e:
        log_error(f"Image gen load failed: {e}")
        return False

def _write_placeholder_image(prompt: str, output_path: str) -> bool:
    if not PIL_AVAILABLE:
        return False
    try:
        img = PILImage.new("RGB", (1024, 1024), color=(16, 24, 40))
        draw = ImageDraw.Draw(img)
        font = ImageFont.load_default()
        draw.text((40, 60), "Image generation unavailable", fill=(255, 255, 255), font=font)
        draw.text((40, 120), "FLUX model could not be loaded.", fill=(200, 200, 200), font=font)
        draw.text((40, 160), f"Prompt: {prompt}", fill=(180, 220, 255), font=font)
        img.save(output_path)
        return True
    except Exception as e:
        log_error(f"Placeholder image failed: {e}")
        return False

def generate_image(prompt: str, output_filename: str = "generated.png",
                    width: int = 1024, height: int = 1024, steps: int = 4, seed: int = 42) -> str:
    router.unload(ModelTier.LIGHT)   # free memory for FLUX; will be reloaded on next chat/vision call
    out_path = str(OUTPUT_DIR / output_filename)
    if not _load_imgen():
        return out_path if _write_placeholder_image(prompt, out_path) else ""
    try:
        image = _imgen_model.generate_image(seed=seed, prompt=prompt,
                                             num_inference_steps=steps, width=width, height=height)
        image.save(path=out_path)
        return out_path
    except Exception as e:
        log_error(f"Image generation failed: {e}")
        return out_path if _write_placeholder_image(prompt, out_path) else ""

def understand_image(image_path: str, prompt: str = "Describe this image in detail.") -> str:
    if not os.path.exists(image_path):
        return f"[ERROR] File not found: {image_path}"
    return _vlm_raw(prompt, [image_path])

def _pdf_extract_text(doc) -> str:
    parts = []
    for i, page in enumerate(doc):
        text = page.get_text("text").strip()
        if text:
            parts.append(f"[Page {i + 1}]\n{text}")
    return "\n\n".join(parts)

def _pdf_pages_to_images(doc) -> list[str]:
    paths = []
    mat = fitz.Matrix(150 / 72, 150 / 72)
    for i, page in enumerate(doc):
        pix = page.get_pixmap(matrix=mat)
        tmp = tempfile.NamedTemporaryFile(suffix=f"_page{i}.png", delete=False, dir=OUTPUT_DIR)
        pix.save(tmp.name)
        tmp.close()
        paths.append(tmp.name)
    return paths

def read_pdf(pdf_path: str, prompt: str = "Summarise the content of this document.") -> str:
    if not PYMUPDF_AVAILABLE:
        return "[ERROR] pymupdf not installed."
    if not os.path.exists(pdf_path):
        return f"[ERROR] File not found: {pdf_path}"
    tmp_paths: list[str] = []
    try:
        doc = fitz.open(pdf_path)
        total_chars = sum(len(p.get_text("text")) for p in doc)
        if total_chars >= 100:
            text = _pdf_extract_text(doc)
            doc.close()
            return _vlm_raw(f"{prompt}\n\n---\nDOCUMENT TEXT:\n{text}", None)
        tmp_paths = _pdf_pages_to_images(doc)
        doc.close()
        return _vlm_raw(prompt, tmp_paths) if tmp_paths else "[ERROR] Could not render PDF pages."
    except Exception as e:
        return f"[ERROR] PDF read failed: {e}"
    finally:
        for p in tmp_paths:
            try: os.unlink(p)
            except OSError: pass

def generate_pdf(title: str, body: str, output_filename: str = "generated.pdf") -> str:
    if not REPORTLAB_AVAILABLE:
        return ""
    out_path = str(OUTPUT_DIR / output_filename)
    try:
        doc = SimpleDocTemplate(out_path, pagesize=A4, leftMargin=2.5 * cm, rightMargin=2.5 * cm,
                                 topMargin=2.5 * cm, bottomMargin=2.5 * cm)
        styles = getSampleStyleSheet()
        story = [Paragraph(title, styles["Title"]), Spacer(1, 0.5 * cm)]
        for para in body.strip().split("\n\n"):
            if para.strip():
                story.append(Paragraph(para.strip(), styles["BodyText"]))
                story.append(Spacer(1, 0.3 * cm))
        doc.build(story)
        return out_path
    except Exception as e:
        log_error(f"PDF generation failed: {e}")
        return ""

def read_docx(docx_path: str, prompt: str = "Summarise the content of this document.") -> str:
    if not DOCX_AVAILABLE:
        return "[ERROR] python-docx not installed."
    if not os.path.exists(docx_path):
        return f"[ERROR] File not found: {docx_path}"
    try:
        doc = DocxDocument(docx_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        tables_text = [" | ".join(c.text.strip() for c in row.cells if c.text.strip())
                        for table in doc.tables for row in table.rows]
        text_body = "\n".join(paragraphs)
        if tables_text:
            text_body += "\n\nTABLES:\n" + "\n".join(t for t in tables_text if t)
        return _vlm_raw(f"{prompt}\n\nDOCUMENT TEXT:\n{text_body}", None)
    except Exception as e:
        return f"[ERROR] DOCX read failed: {e}"

def generate_docx(title: str, body: str, output_filename: str = "generated.docx") -> str:
    if not DOCX_AVAILABLE:
        return ""
    out_path = str(OUTPUT_DIR / output_filename)
    try:
        doc = DocxDocument()
        doc.add_heading(title, level=0)
        for para in body.strip().split("\n\n"):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.runs[0].font.size = Pt(11)
        doc.save(out_path)
        return out_path
    except Exception as e:
        log_error(f"DOCX generation failed: {e}")
        return ""


# -----------
# VOICE GATE
# -----------

voice_gate = None
if ENABLE_VOICE_GATE and VOICEGATE_IMPORTABLE:
    voice_gate = VoiceGate(voiceprint_path=VOICEPRINT_PATH, threshold=VOICE_GATE_THRESHOLD)
    if voice_gate.voiceprint is None:
        log_error(f"Voice gate enabled but no voiceprint at {VOICEPRINT_PATH}. "
                   f"Enroll first: python A14_VoiceGate_session.py enroll")
        voice_gate = None

def _record_gate_clip(path: str, seconds: float = 3.0) -> None:
    if not SOUNDFILE_AVAILABLE:
        raise RuntimeError("soundfile not installed; can't capture a gate clip.")
    audio = sd.rec(int(seconds * SAMPLE_RATE), samplerate=SAMPLE_RATE, channels=1,
                   dtype="float32", device=0)
    sd.wait()
    sf.write(path, audio, SAMPLE_RATE)

def unlock_session() -> bool:
    if voice_gate is None:
        return True
    path = str(OUTPUT_DIR / "startup_verify.wav")
    speak_interruptible("Please say something so I can verify it's you.")
    _record_gate_clip(path)
    score, passed = voice_gate.verify_file(path)
    log_info(f"Startup verification score={score:.3f} passed={passed}")
    return passed

def guarded_action(action_name: str, action_fn, *args, **kwargs):
    if voice_gate is None:
        return action_fn(*args, **kwargs)
    path = str(OUTPUT_DIR / "confirm_verify.wav")
    print(f"'{action_name}' requires re-verification.")
    _record_gate_clip(path, seconds=3.0)
    score, passed = voice_gate.verify_file(path)
    log_info(f"Re-verification for '{action_name}' score={score:.3f} passed={passed}")
    if not passed:
        return f"Blocked: '{action_name}' requires your verified voice."
    return action_fn(*args, **kwargs)


# -------------
# APPS / TOOLS
# -------------

def get_installed_apps() -> list[str]:
    app_dirs = ["/Applications", "/System/Applications", "/System/Applications/Utilities",
                os.path.expanduser("~/Applications")]
    apps = []
    for d in app_dirs:
        if os.path.exists(d):
            apps += [f.replace(".app", "") for f in os.listdir(d) if f.endswith(".app")]
    return apps

BROWSER_APPS = {"Google Chrome", "Safari", "Firefox", "Arc", "Brave Browser"}
MUSIC_APPS = {"Music"}
installed = set(get_installed_apps())
browsers = installed & BROWSER_APPS
music_apps = installed & MUSIC_APPS

tools: dict[str, dict] = {
    "open_app": {"description": "Open any installed application", "args": {"app": "name of any installed application"}},
    "quit_app": {"description": "Force quit an application", "args": {"app": "name of the application to quit"}},
    "web_search": {
        "description": "Search the web and return an accurate answer. Use for current events, prices, scores, or anything time-sensitive.",
        "args": {"query": "search query string"},
    },
    "update_memory": {
        "description": "Explicitly store or update a fact about the user in long-term memory.",
        "args": {"key": "fact key in snake_case", "value": "fact value"},
    },
    "generate_image": {
        "description": "Generate an image from a text prompt.",
        "args": {"prompt": "description of the image", "output_filename": "e.g. generated.png"},
    },
    "understand_image": {
        "description": "Describe or answer a question about an existing image file.",
        "args": {"image_path": "path to the image", "prompt": "what to ask about the image"},
    },
    "generate_pdf": {
        "description": "Create a PDF document from a title and body text.",
        "args": {"title": "document title", "body": "document body (use \\n\\n between paragraphs)"},
    },
    "read_pdf": {
        "description": "Read and summarise or answer questions about an existing PDF file.",
        "args": {"pdf_path": "path to the PDF", "prompt": "what to ask about the document"},
    },
    "generate_docx": {
        "description": "Create a Word document from a title and body text.",
        "args": {"title": "document title", "body": "document body (use \\n\\n between paragraphs)"},
    },
    "read_docx": {
        "description": "Read and summarise or answer questions about an existing Word document.",
        "args": {"docx_path": "path to the .docx file", "prompt": "what to ask about the document"},
    },
}

if browsers:
    tools["search_web"] = {"description": "Open a URL or search query in a browser",
                            "args": {"query": "search term or full URL", "browser": f"one of: {browsers}"}}

if music_apps:
    tools["play_music"] = {"description": "Play a song or artist in Apple Music.",
                            "args": {"query": "full search string e.g. 'song name by artist'",
                                      "app": f"one of: {music_apps}", "shuffle": "true or false (optional)"}}
    tools["pause_music"] = {"description": "Pause or resume currently playing music", "args": {"app": f"one of: {music_apps}"}}
    tools["shuffle_music"] = {"description": "Toggle shuffle on or off", "args": {"enabled": "true or false", "app": f"one of: {music_apps}"}}
    tools["repeat_music"] = {"description": "Set repeat mode", "args": {"mode": "one of: one, all, off", "app": f"one of: {music_apps}"}}
    tools["play_library"] = {"description": "Play a specific album or playlist from your local Music library.",
                              "args": {"type": "one of: album, playlist", "query": "album or playlist name",
                                        "app": f"one of: {music_apps}", "shuffle": "true or false (optional)"}}


def run_applescript(script: str) -> str:
    result = subprocess.run(["osascript", "-e", script], capture_output=True, text=True, check=False)
    if result.returncode != 0:
        raise RuntimeError(result.stderr.strip() or "AppleScript failed")
    return result.stdout.strip()


def _do_quit_app(app: str) -> str:
    run_applescript(f'tell application "{app}" to quit')
    return f"Quit {app}."

def _do_update_memory(key: str, value: str) -> str:
    if key and value:
        set_fact(key, value)
        store_semantic_memory(f"User fact — {key}: {value}", {"type": "user_fact", "key": key})
        return f"Got it, I'll remember that {key.replace('_', ' ')}: {value}."
    return "Couldn't store that — missing key or value."


def execute_tool(name: str, args: dict, user_input: str = "") -> str:
    name = name.replace(" ", "_").lower()
    try:
        if name == "open_app":
            app = args.get("app", "").strip()
            if app not in installed:
                return f"App '{app}' not found."
            subprocess.Popen(["open", "-a", app])
            return f"Opened {app}."

        elif name == "quit_app":
            app = args.get("app", "").strip()
            return guarded_action("quit_app", _do_quit_app, app)

        elif name == "search_web":
            query = args.get("query", "").strip()
            browser = args.get("browser", next(iter(browsers))).strip()
            if not query.startswith("http"):
                query = "https://www.google.com/search?q=" + query.replace(" ", "+")
            subprocess.Popen(["open", "-a", browser, query])
            return f"Opened '{query}' in {browser}."

        elif name == "web_search":
            query = args.get("query", "").strip()
            force_fresh = any(w in user_input.lower() for w in ["latest", "again", "refresh", "new search"])
            return do_web_search(query, user_input, force_fresh=force_fresh)

        elif name == "update_memory":
            key = args.get("key", "").strip()
            value = args.get("value", "").strip()
            return guarded_action("update_memory", _do_update_memory, key, value)

        elif name == "generate_image":
            path = generate_image(args.get("prompt", "").strip(), args.get("output_filename", "generated.png"))
            return f"Image saved to {path}." if path else "Image generation failed."

        elif name == "understand_image":
            return understand_image(args.get("image_path", "").strip(), args.get("prompt", "Describe this image in detail."))

        elif name == "generate_pdf":
            path = generate_pdf(args.get("title", "Untitled"), args.get("body", ""), args.get("output_filename", "generated.pdf"))
            return f"PDF saved to {path}." if path else "PDF generation failed."

        elif name == "read_pdf":
            return read_pdf(args.get("pdf_path", "").strip(), args.get("prompt", "Summarise the content of this document."))

        elif name == "generate_docx":
            path = generate_docx(args.get("title", "Untitled"), args.get("body", ""), args.get("output_filename", "generated.docx"))
            return f"Word document saved to {path}." if path else "DOCX generation failed."

        elif name == "read_docx":
            return read_docx(args.get("docx_path", "").strip(), args.get("prompt", "Summarise the content of this document."))

        elif name == "play_music":
            query = args.get("query", "").strip()
            if not query:
                return "No query specified for play_music."
            shuffle = str(args.get("shuffle", "false")).strip().lower() == "true"
            if shuffle:
                run_applescript('tell application "Music" to set shuffle enabled to true')
            result = subprocess.run(["shortcuts", "run", "Play Apple Music", "--input-string", query],
                                     capture_output=True, text=True)
            if result.returncode != 0:
                subprocess.Popen(["open", f"music://music.apple.com/search?term={quote(query)}"])
                return f"Couldn't auto-play '{query}'. Opened Music search instead."
            return f"Playing '{query}' on Music{' with shuffle' if shuffle else ''}."

        elif name == "pause_music":
            app = args.get("app", next(iter(music_apps))).strip()
            run_applescript(f'tell application "{app}" to pause')
            return f"Paused {app}."

        elif name == "shuffle_music":
            enabled = args.get("enabled", "true").strip().lower()
            app = args.get("app", next(iter(music_apps))).strip()
            value = "true" if enabled == "true" else "false"
            run_applescript(f'tell application "{app}" to set shuffle enabled to {value}')
            return f"Shuffle {'enabled' if value == 'true' else 'disabled'} on {app}."

        elif name == "repeat_music":
            mode = args.get("mode", "all").strip().lower()
            app = args.get("app", next(iter(music_apps))).strip()
            applescript_mode = {"one": "one", "all": "all", "off": "off"}.get(mode, "off")
            run_applescript(f'tell application "{app}" to set song repeat to {applescript_mode}')
            return f"Repeat set to '{applescript_mode}' on {app}."

        elif name == "play_library":
            ptype = args.get("type", "playlist").strip().lower()
            query = args.get("query", "").strip()
            app = args.get("app", next(iter(music_apps))).strip()
            shuffle = str(args.get("shuffle", "false")).strip().lower() == "true"
            run_applescript('tell application "Music" to activate')
            time.sleep(2)
            if shuffle:
                run_applescript('tell application "Music" to set shuffle enabled to true')
            if ptype == "album":
                run_applescript(f'''tell application "Music"
                    activate
                    set sr to (every track of playlist "Library" whose album is "{query}")
                    if sr is not {{}} then play item 1 of sr
                end tell''')
            elif ptype == "playlist":
                run_applescript(f'''tell application "Music"
                    activate
                    play playlist "{query}"
                end tell''')
            else:
                return f"Unknown library type: {ptype}. Use album or playlist."
            return f"Playing {ptype}{' — ' + query if query else ''} on {app}{' with shuffle' if shuffle else ''}."

        else:
            return f"Unknown tool: {name}"

    except Exception as e:
        return f"Error executing {name}: {str(e)}"


# --------------
# SYSTEM PROMPT
# --------------

SYSTEM_PROMPT_TEMPLATE = """Today's date is {date}. Current time is {time}.
You are Gwen, a precise and practical AI assistant running primarily on-device. You are currently answering in {tier_name} mode, selected manually by the user — you do not choose or change this yourself.

{global_context}

{semantic_context}

CORE BEHAVIOUR — LOW HALLUCINATION:
- Give only the final answer. No hidden reasoning, no chain-of-thought, no <think> blocks, no scratchpad, no self-correction narrated out loud.
- Do not explain how you arrived at an answer unless explicitly asked.
- Do not invent facts, sources, statistics, or quotes. If you are not sure, say so plainly instead of guessing.
- Never claim to have checked, searched, or verified something unless a tool result actually appears in this conversation.
- Answer in 1-3 sentences unless the user explicitly asks for a list, code, or more detail.

TOOLS:
You have access to tools, but ONLY output a JSON tool call when the user is explicitly requesting a system action, a file operation, or a fact that needs live/current information.
NEVER output a JSON tool call for greetings, casual conversation, or thank-you messages.

When a tool call IS appropriate, output exactly one raw JSON object and nothing else:
{{"name": "tool_name", "arguments": {{"key": "value"}}}}

Available tools:
{tools}

Tool names must be exactly: {tool_names}.

STRICT RULE ON THE INTERNET — follow without exception:
- For ANY question about people, birthdays, dates, current facts, news, prices, scores, or live events — ALWAYS call web_search. Never answer from memory when currency or accuracy is uncertain.
- If you are not 100 percent certain of a fact, call web_search instead of guessing.
- If you start to feel unsure mid-answer, stop and call web_search instead of continuing to generate text.

Examples of tool calls:
- "play Blinding Lights by The Weeknd" -> {{"name": "play_music", "arguments": {{"query": "Blinding Lights by The Weeknd", "app": "Music", "shuffle": "false"}}}}
- "who is the current prime minister of the UK" -> {{"name": "web_search", "arguments": {{"query": "current UK prime minister"}}}}
- "remember I stopped working on project X" -> {{"name": "update_memory", "arguments": {{"key": "stopped_projects", "value": "project X"}}}}
- "summarise this PDF at ~/Downloads/report.pdf" -> {{"name": "read_pdf", "arguments": {{"pdf_path": "~/Downloads/report.pdf", "prompt": "Summarise this document."}}}}

Examples of plain text (NO tool call):
- "hey how are you" -> respond conversationally
- "what is the capital of France" -> answer directly (stable, non-time-sensitive fact)"""


def build_system_prompt(global_context: str, semantic_context: list[str], tier: ModelTier) -> str:
    return SYSTEM_PROMPT_TEMPLATE.format(
        date=datetime.now().strftime("%B %d, %Y"),
        time=datetime.now().strftime("%I:%M %p"),
        tier_name=tier.value,
        global_context=global_context if global_context else "",
        semantic_context=f"[Relevant memories]\n{chr(10).join(semantic_context)}" if semantic_context else "",
        tools=json.dumps(tools, indent=2),
        tool_names=", ".join(tools.keys()),
    )


def parse_tool_call(text: str) -> Optional[dict]:
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL)
    start = text.find("{")
    if start == -1:
        return None
    depth = 0
    for i, ch in enumerate(text[start:], start):
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                try:
                    data = json.loads(text[start:i + 1])
                    return data if "name" in data and "arguments" in data else None
                except json.JSONDecodeError:
                    return None
    return None


# -------------------------
# LATENCY / MEMORY LOGGING
# -------------------------

def log_latency(query: str, model_tier: str, timings: dict):
    conn = get_db()
    conn.execute("""INSERT INTO latency_log
        (session_id, query_preview, complexity, model_tier, stt_ms, memory_ms, inference_ms, tts_ms, total_ms, logged_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
        (SESSION_ID, query[:80], "unknown", model_tier, timings.get("stt_ms"), timings.get("memory_ms"),
         timings.get("inference_ms"), timings.get("tts_ms"), timings.get("total_ms"), datetime.now().isoformat()))
    conn.commit()
    conn.close()

def log_memory_retrieval(query: str, facts: dict, semantic_hits: list[str]):
    conn = get_db()
    conn.execute("""INSERT INTO memory_retrieval_log
        (session_id, query_preview, facts_injected, semantic_hits, logged_at) VALUES (?, ?, ?, ?, ?)""",
        (SESSION_ID, query[:80], json.dumps(facts), json.dumps(semantic_hits), datetime.now().isoformat()))
    conn.commit()
    conn.close()


# ----------------
# SESSION SUMMARY
# ----------------

chat_history: list[dict] = []

def summarize_and_save_session():
    if not chat_history or not SMALL_MODEL_AVAILABLE:
        return
    history_text = "\n".join(f"{m['role'].capitalize()}: {m['content'][:200]}" for m in chat_history[-20:])
    prompt = f"""Summarize this conversation in 2-3 sentences. Focus on what the user asked, what was accomplished, and any important facts mentioned.

Conversation:
{history_text}

Summary:"""
    try:
        response = generate_lm(small_model, small_tokenizer,
                                prompt=f"<|im_start|>user\n{prompt}<|im_end|>\n<|im_start|>assistant\n",
                                max_tokens=150, sampler=make_sampler(temp=0.3), verbose=False)
        summary = strip_special_tokens(clean_response(extract_mlx_text(response), "qwen"))
        conn = get_db()
        conn.execute("INSERT INTO session_summaries (summary, created_at) VALUES (?, ?)",
                     (summary, datetime.now().isoformat()))
        conn.commit()
        conn.close()
        store_semantic_memory(f"Session summary ({datetime.now().strftime('%Y-%m-%d')}): {summary}", {"type": "session_summary"})
        log_info(f"Saved session summary: {summary[:80]}...")
    except Exception as e:
        log_error(f"Summary failed: {e}")


# ---------------------
# MAIN QUERY PIPELINE
# ---------------------

model_ctrl: ModelModeController
input_ctrl: InputModeController


def process_query(transcription: str, stt_ms: float) -> Optional[bool]:
    global deep_search_mode

    transcription = transcription.strip()
    if not transcription:
        return False

    t_start = time.time()
    timings = {"stt_ms": stt_ms}
    print(f"[{_ts()}] You: {transcription}")

    if is_exit(transcription):
        speak_interruptible("Shutting down. Goodbye.")
        beep_deactivate()
        summarize_and_save_session()
        return None

    if is_deactivation(transcription):
        speak_interruptible("Going to sleep.")
        beep_deactivate()
        summarize_and_save_session()
        chat_history.clear()
        return False

    if "toggle deep search" in transcription.lower():
        deep_search_mode = not deep_search_mode
        msg = f"Deep search {'enabled' if deep_search_mode else 'disabled'}."
        print(f"[System] {msg}")
        speak_interruptible(msg)
        return True

    if input_ctrl.process(transcription):
        return True

    if model_ctrl.is_status_query(transcription):
        speak_interruptible(model_ctrl.status_text())
        return True

    switch_ack = model_ctrl.process(transcription)
    if switch_ack is not None:
        print(f"[System] {switch_ack}")
        speak_interruptible(switch_ack)
        return True

    detect_and_apply_fact_updates(transcription)

    t_mem = time.time()
    facts = get_all_facts()
    global_context = build_global_context()
    semantic_hits = retrieve_semantic_memory(transcription, n_results=3)
    timings["memory_ms"] = (time.time() - t_mem) * 1000
    log_memory_retrieval(transcription, facts, semantic_hits)

    system_prompt = build_system_prompt(global_context, semantic_hits, model_ctrl.tier)

    force_search = any(re.search(p, transcription.lower()) for p in FORCE_SEARCH_PATTERNS)
    if force_search:
        log_info("Force search triggered.")
        final_response = sanitize_response(do_web_search(transcription, transcription))
    else:
        tier_before = model_ctrl.tier
        t_inf = time.time()
        raw = router.chat(system_prompt=system_prompt, history=chat_history[-10:], user_input=transcription)
        timings["inference_ms"] = (time.time() - t_inf) * 1000

        if model_ctrl.tier != tier_before:
            note = f"Heads up — {tier_before.value} mode hit its limit, I've switched to {model_ctrl.tier.value} mode."
            print(f"[System] {note}")
            if input_ctrl.is_voice:
                speak_interruptible(note)

        tool_call = parse_tool_call(raw)
        if tool_call:
            final_response = execute_tool(tool_call["name"], tool_call["arguments"], transcription)
            print(f"[Tool] {tool_call['name']}: {final_response}")
        else:
            final_response = sanitize_response(raw)

    print(f"[{_ts()}] Gwen: {final_response}")

    t_tts = time.time()
    if input_ctrl.is_voice:
        try:
            speak_interruptible(final_response)
        except Exception as e:
            log_error(f"speak failed: {e}")
            sd.stop()
    timings["tts_ms"] = (time.time() - t_tts) * 1000
    timings["total_ms"] = (time.time() - t_start) * 1000

    print(f"[Latency] STT:{timings.get('stt_ms', 0):.0f}ms | Mem:{timings.get('memory_ms', 0):.0f}ms | "
          f"Inf:{timings.get('inference_ms', 0):.0f}ms | TTS:{timings.get('tts_ms', 0):.0f}ms | "
          f"Total:{timings.get('total_ms', 0) / 1000:.2f}s | Tier:{model_ctrl.tier.value}")

    log_latency(transcription, model_ctrl.tier.value, timings)
    chat_history.append({"role": "user", "content": transcription})
    chat_history.append({"role": "assistant", "content": final_response})
    return True


# ----------
# MAIN LOOP
# ----------

def start_chat():
    global model_ctrl, input_ctrl

    startup_tier = ModelTier[DEFAULT_MODEL_TIER_NAME]
    model_ctrl = ModelModeController(router, startup_tier)
    input_ctrl = InputModeController(DEFAULT_INPUT_MODE)

    print(f"\n[Super Gwen] Session: {SESSION_ID}")
    print(f"[Memory] DB: {DB_PATH}")
    print(f"[Memory] ChromaDB: {'enabled' if CHROMA_AVAILABLE else 'disabled'}")
    print(f"[Model] Starting tier: {model_ctrl.tier.value}  (say 'switch to max mode' etc. to change)")
    print(f"[Voice Gate] {'enabled' if voice_gate else 'disabled'}\n")

    if not unlock_session():
        print("[Gwen] Session not started: voice not recognized.")
        return

    while True:
        if input_ctrl.is_voice:
            listen_for_wake_word()
            beep_activate()
            speak_interruptible(pick_activation_phrase())
            time.sleep(0.2)
            print("[Gwen] Activated.")
        else:
            print("[Gwen] Text mode — type your message ('exit' to quit).")

        while True:
            if input_ctrl.is_voice:
                transcription, stt_ms = listen_for_prompt()
                if transcription is None:
                    print("[Gwen] Timeout — going dormant.")
                    beep_timeout()
                    break
            else:
                transcription = input("You (text): ").strip()
                stt_ms = 0.0

            result = process_query(transcription, stt_ms)

            if result is None:
                return
            if not result:
                break
            if input_ctrl.is_voice:
                beep_activate()


if __name__ == "__main__":
    start_chat()